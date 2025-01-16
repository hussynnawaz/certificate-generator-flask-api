import os
import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document
from docx2pdf import convert
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.text import MIMEText
import re

# Firebase setup
cred = credentials.Certificate("firebaseConfig.json")  # Use your Firebase service account JSON file
firebase_admin.initialize_app(cred)
db = firestore.client()

# Send email function
def send_email(to_email, subject, body, attachment_path):
    from_email = "sardar.hussyn@gmail.com"
    app_password = "axmm aayj djsd cqql"  # Use your generated App Password here

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))
    
    with open(attachment_path, 'rb') as f:
        attach = MIMEApplication(f.read(), _subtype="pdf")
        attach.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment_path))
        msg.attach(attach)

    try:
        # Set up the SMTP server
        server = smtplib.SMTP_SSL('smtp.gmail.com', 465)  # Using Gmail's SMTP server
        server.set_debuglevel(1)  # Enable debugging to get detailed output
        server.login(from_email, app_password)  # Use the generated App Password here
        server.sendmail(from_email, to_email, msg.as_string())
        server.quit()
        print(f"Email sent successfully to {to_email}")
    except smtplib.SMTPAuthenticationError as e:
        print(f"Authentication failed: {str(e)}")  # Improved error message for authentication issues
    except Exception as e:
        print(f"Error: {str(e)}")

# Replace placeholder in certificate
def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)

def replace_info(doc, name, string):
    reg = re.compile(r"" + string)
    replace = r"" + name
    docx_replace_regex(doc, reg, replace)

def replace_participant_name(doc, name):
    string = "{Name Surname}"  # Placeholder for name
    replace_info(doc, name, string)

# Function to generate certificate and send via email
def generate_and_send_certificate(user_id):
    # Get user data from Firebase
    user_ref = db.collection('users').document(user_id)
    user = user_ref.get()

    if user.exists:
        user_data = user.to_dict()
        name = user_data.get("name")
        email = user_data.get("email")

        # Ensure directories exist before saving files
        doc_output_dir = 'Output/Doc'
        pdf_output_dir = 'Output/PDF'
        os.makedirs(doc_output_dir, exist_ok=True)
        os.makedirs(pdf_output_dir, exist_ok=True)

        # Generate certificate document
        certificate_file = "Certificate.docx"  # Ensure this template file exists
        doc = Document(certificate_file)
        replace_participant_name(doc, name)  # Replace name in the template

        # Save the document locally as DOCX
        output_docx_path = os.path.join(doc_output_dir, f'{name}.docx')
        doc.save(output_docx_path)

        # Convert to PDF
        pdf_path = os.path.join(pdf_output_dir, f'{name}.pdf')
        convert(output_docx_path, pdf_path)

        # Save PDF locally (ensure output directories exist)
        local_pdf_path = f'./local_pdfs/{name}.pdf'
        os.makedirs(os.path.dirname(local_pdf_path), exist_ok=True)
        convert(output_docx_path, local_pdf_path)

        # Send email with the PDF attachment
        send_email(email, "Your Certificate", "Attached is your certificate.", local_pdf_path)
        return {"status": "success", "message": f"Certificate for {name} sent to {email}"}
    else:
        return {"status": "error", "message": f"User {user_id} not found in Firebase."}

# Example usage: Pass in a user ID
result = generate_and_send_certificate("5RaQo3NinhWCPOkSPDC779ceKOq1")
print(result)
